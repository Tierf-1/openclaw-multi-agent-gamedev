"""
OpenClaw Web 后端 - FastAPI REST API 服务

提供企业级仪表盘后端:
- 系统概览 API
- Agent 管理 API
- 流水线管理 API (含异步执行)
- 消息队列 API
- 沙盒管理 API
- 实时日志 API
- 规则资产 API
- 智能体模型配置 API
"""

import json
import sys
import os
import asyncio
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

from fastapi import FastAPI, HTTPException, Query, BackgroundTasks, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel

logger = logging.getLogger(__name__)

# 确保项目根目录在路径中
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT / "src"))


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Pydantic 请求/响应模型
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class RequirementRequest(BaseModel):
    """需求处理请求"""
    user_input: str
    req_type: Optional[str] = None
    req_scale: Optional[str] = None


class AgentConfigRequest(BaseModel):
    """智能体模型配置请求"""
    provider: Optional[str] = None
    model: Optional[str] = None
    api_key: Optional[str] = None
    base_url: Optional[str] = None
    temperature: Optional[float] = None
    max_tokens: Optional[int] = None
    enabled: Optional[bool] = None


class MessageRequest(BaseModel):
    """消息发送请求"""
    from_agent: str
    to_agent: str
    msg_type: str = "handoff"
    payload: Dict[str, Any] = {}


class RenameRequest(BaseModel):
    """流水线重命名请求"""
    name: str


class CustomRuleRequest(BaseModel):
    """自定义规则模板请求"""
    content: str
    filename: Optional[str] = None


class NewAgentRequest(BaseModel):
    """新智能体创建请求"""
    agent_name: str
    agent_icon: str = "🤖"
    role: str = ""
    persona: str = ""
    group: str = "implementation"
    entry_content: str = ""
    sandbox: bool = True


class AgentPluginRequest(BaseModel):
    """智能体插件配置"""
    plugin_id: str
    enabled: bool = True
    config: Dict[str, Any] = {}


class AgentMCPRequest(BaseModel):
    """智能体MCP服务器配置"""
    server_name: str
    server_url: str = ""
    enabled: bool = True
    tools: List[str] = []


class AgentIntegrationRequest(BaseModel):
    """智能体集成配置"""
    integration_id: str
    enabled: bool = True
    config: Dict[str, Any] = {}


class AgentSkillBindRequest(BaseModel):
    """绑定/解绑技能到智能体"""
    skill_id: str
    enabled: bool = True


class AgentMemoryRequest(BaseModel):
    """智能体记忆条目"""
    title: str
    content: str
    memory_type: str = "knowledge"  # knowledge / preference / context


class ApiResponse(BaseModel):
    """通用API响应"""
    success: bool = True
    message: str = ""
    data: Any = None
    timestamp: str = ""

    def __init__(self, **kwargs):
        if "timestamp" not in kwargs:
            kwargs["timestamp"] = datetime.now().isoformat()
        super().__init__(**kwargs)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 全局状态（进程内单例）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class AppState:
    """应用全局状态"""
    def __init__(self, project_root: str):
        self.project_root = Path(project_root)
        self.orchestrator = None
        self.rule_loader = None
        self.codebuddy_adapter = None
        self.model_config_mgr = None
        self.llm_invoker = None
        self.inventory = None
        self.initialized = False
        # 异步执行中的流水线
        self._running_pipelines: Dict[str, dict] = {}

    def initialize(self):
        """延迟初始化"""
        if self.initialized:
            return

        # ── 初始化 SQLite 数据库 ──────────────────────────────
        try:
            from core.database import init_db
            db_path = self.project_root / "data" / "openclaw.db"
            init_db(db_path)
            logger.info(f"数据库已初始化: {db_path}")
        except Exception as e:
            logger.warning(f"数据库初始化失败 (将降级为纯内存模式): {e}")

        from core.orchestrator import Orchestrator
        from adapters.rule_loader import RuleLoader
        from adapters.codebuddy_adapter import CodeBuddyAdapter
        from utils.file_ops import safe_yaml_read
        from core.llm_adapter import ModelConfigManager, LLMInvoker

        # 读取配置
        config_path = self.project_root / "config" / "system.yaml"
        config = safe_yaml_read(config_path) or {}
        parallel_config = config.get("parallel_architecture", {})

        # 初始化模型配置管理器
        model_config_path = self.project_root / "config" / "agent_models.json"
        self.model_config_mgr = ModelConfigManager(str(model_config_path))
        self.llm_invoker = LLMInvoker(self.model_config_mgr)

        # 初始化编排器
        self.orchestrator = Orchestrator(
            project_root=str(self.project_root),
            config={
                "sandbox_root": parallel_config.get("sandbox_root", ".sandboxes"),
                "protected_paths": ["rules"],
                "max_parallel_agents": parallel_config.get("max_parallel_agents", 4)
            }
        )

        # 初始化规则加载器
        self.rule_loader = RuleLoader(str(self.project_root / "rules"))
        self.inventory = self.rule_loader.scan_all()

        # 初始化CodeBuddy适配器
        self.codebuddy_adapter = CodeBuddyAdapter(str(self.project_root))
        self.codebuddy_adapter.initialize(rule_loader=self.rule_loader)

        # 系统初始化
        self.orchestrator.setup()

        # 初始化Agent沙盒 + 注册Agent实例
        from core.sandbox import SandboxConfig
        agents_config = config.get("agents", {})
        for agent_key, agent_def in agents_config.items():
            agent_id = agent_def.get("id", agent_key)
            if agent_def.get("sandbox", False):
                sc = SandboxConfig(
                    agent_id=agent_id,
                    agent_name=agent_def.get("name", agent_key),
                    sandbox_root=parallel_config.get("sandbox_root", ".sandboxes"),
                    read_permissions=["rules/**"],
                    write_permissions=[".GameDev/**"]
                )
                self.orchestrator.sandbox_mgr.create_sandbox(sc)

        # 注册所有Agent实例到编排器
        self._register_agents()

        # ── 从数据库恢复 Pipeline 内存状态 ────────────────────
        self._restore_pipelines_from_db()

        self.initialized = True

    def _register_agents(self):
        """注册所有Agent实例到编排器"""
        try:
            from agents import AGENT_REGISTRY
            for agent_id, agent_cls in AGENT_REGISTRY.items():
                try:
                    agent_instance = agent_cls(
                        agent_id=agent_id,
                        sandbox_mgr=self.orchestrator.sandbox_mgr,
                        message_queue=self.orchestrator.mq,
                        context_mgr=self.orchestrator.ctx_mgr,
                    )
                    self.orchestrator.register_agent(agent_id, agent_instance)
                    logger.info(f"注册Agent: {agent_id} ({agent_cls.__name__})")
                except Exception as e:
                    logger.warning(f"注册Agent {agent_id} 失败: {e}, 将使用模拟模式")
        except Exception as e:
            logger.warning(f"加载Agent模块失败: {e}, 所有Agent将使用模拟模式")

    def _restore_pipelines_from_db(self):
        """从数据库恢复 Pipeline 到内存引擎（服务重启后保持数据）"""
        try:
            from core.database import load_all_pipelines
            from core.pipeline import PipelineInstance, PipelineStep, ReqType, ReqScale

            pipelines = load_all_pipelines()
            restored = 0
            for p in pipelines:
                pid = p["pipeline_id"]
                # 跳过已存在的（正常情况下重启后内存为空）
                if self.orchestrator.pipeline_engine.get_pipeline(pid):
                    continue
                try:
                    req_type = ReqType(p.get("req_type", "FEATURE"))
                    req_scale = ReqScale(p.get("req_scale", "M"))
                except (ValueError, KeyError):
                    req_type = ReqType.FEATURE
                    req_scale = ReqScale.M

                steps = [
                    PipelineStep(
                        stage=s.get("stage", s.get("name", "")),
                        agent_id=s.get("agent_id", ""),
                        execution=s.get("execution", "sequential"),
                        parallel_with=s.get("parallel_with", []),
                        quality_gate=s.get("quality_gate"),
                        optional=bool(s.get("optional", False)),
                        status=s.get("status", "pending"),
                    )
                    for s in (p.get("steps") or [])
                ]

                instance = PipelineInstance(
                    pipeline_id=pid,
                    req_id=p.get("req_id", ""),
                    req_type=req_type.value,
                    req_scale=req_scale.value,
                    req_name=p.get("req_name", ""),
                    steps=steps,
                    current_step_index=p.get("current_step_index", 0),
                    status=p.get("status", "created"),
                    bug_rounds=p.get("bug_rounds", 0),
                    created_at=p.get("created_at", ""),
                    started_at=p.get("started_at"),
                    completed_at=p.get("completed_at"),
                )
                self.orchestrator.pipeline_engine._instances[pid] = instance
                restored += 1

            if restored:
                logger.info(f"从数据库恢复了 {restored} 条 Pipeline")
        except Exception as e:
            logger.warning(f"从数据库恢复 Pipeline 失败: {e}")

    async def execute_pipeline_async(self, pipeline_id: str, user_input: str):
        """异步执行流水线"""
        try:
            self._running_pipelines[pipeline_id] = {
                "status": "running",
                "started_at": datetime.now().isoformat(),
                "progress": 0,
            }

            # 获取流水线实例
            pipeline = self.orchestrator.pipeline_engine.get_pipeline(pipeline_id)
            if not pipeline:
                self._running_pipelines[pipeline_id]["status"] = "failed"
                self._running_pipelines[pipeline_id]["error"] = "流水线不存在"
                return

            # 设置流水线为运行状态
            pipeline.status = "running"
            pipeline.started_at = datetime.now().isoformat()
            self.orchestrator._active_pipelines[pipeline_id] = pipeline

            # 执行流水线 (在线程池中运行同步方法)
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                self.orchestrator._execute_pipeline,
                pipeline
            )

            self._running_pipelines[pipeline_id] = {
                "status": "completed",
                "completed_at": datetime.now().isoformat(),
                "result": result,
            }
            logger.info(f"流水线 {pipeline_id} 执行完成")

        except Exception as e:
            logger.error(f"流水线 {pipeline_id} 执行失败: {e}")
            self._running_pipelines[pipeline_id] = {
                "status": "failed",
                "error": str(e),
            }
            # 更新pipeline状态
            pipeline = self.orchestrator.pipeline_engine.get_pipeline(pipeline_id)
            if pipeline:
                pipeline.status = "failed"


# 全局状态实例
_app_state: Optional[AppState] = None


def get_state() -> AppState:
    """获取应用状态"""
    global _app_state
    if _app_state is None:
        _app_state = AppState(str(PROJECT_ROOT))
    if not _app_state.initialized:
        _app_state.initialize()
    return _app_state


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# FastAPI 应用工厂
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def create_app(project_root: str = None) -> FastAPI:
    """
    创建FastAPI应用

    Args:
        project_root: 项目根目录

    Returns:
        FastAPI实例
    """
    global _app_state

    if project_root:
        _app_state = AppState(project_root)

    app = FastAPI(
        title="OpenClaw Multi-Agent GameDev Dashboard",
        description="🎮 多智能体并行游戏开发团队系统 - 企业级仪表盘",
        version="1.0.0",
        docs_url="/api/docs",
        redoc_url="/api/redoc"
    )

    # CORS
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    # ── 前端静态文件 ─────────────────────────────────
    # 优先使用 React 前端构建产物 (frontend/dist)
    frontend_dist = PROJECT_ROOT / "frontend" / "dist"
    legacy_static = Path(__file__).parent / "static"

    # 挂载前端资产目录 (CSS/JS chunks)
    if frontend_dist.exists() and (frontend_dist / "assets").exists():
        app.mount("/assets", StaticFiles(directory=str(frontend_dist / "assets")), name="assets")

    # 兼容旧版静态文件
    if legacy_static.exists():
        app.mount("/static", StaticFiles(directory=str(legacy_static)), name="static")

    # ── 系统概览 API ─────────────────────────────────
    @app.get("/api/overview")
    async def get_overview():
        """获取系统概览"""
        state = get_state()

        overview = {
            "system": {
                "name": "OpenClaw Multi-Agent GameDev",
                "version": "1.0.0",
                "status": "running",
                "uptime": datetime.now().isoformat()
            },
            "agents": {
                "total": 8,
                "active": len(state.orchestrator.sandbox_mgr.list_active_sandboxes()),
                "parallel_groups": 5
            },
            "pipelines": {
                "active": len(state.orchestrator._active_pipelines),
                "types": 9,
                "quality_gates": 3
            },
            "inventory": state.inventory.statistics if state.inventory else {},
            "message_queue": state.orchestrator.mq.get_queue_stats()
        }

        return ApiResponse(data=overview)

    # ── Agent 管理 API ───────────────────────────────
    @app.get("/api/agents")
    async def list_agents():
        """获取所有Agent信息（含真实模型配置状态）"""
        state = get_state()

        from utils.file_ops import safe_yaml_read
        config = safe_yaml_read(state.project_root / "config" / "system.yaml") or {}
        agents_config = config.get("agents", {})

        # 获取所有模型配置，用于判定真实状态
        model_configs = {}
        if state.model_config_mgr:
            for cfg in state.model_config_mgr.get_all_configs_masked():
                model_configs[cfg["agent_id"]] = cfg

        agents_list = []
        for agent_key, agent_def in agents_config.items():
            agent_id = agent_def.get("id", agent_key)
            sandbox_path = state.orchestrator.sandbox_mgr.get_sandbox_path(agent_id)

            # 获取上下文信息
            ctx = state.orchestrator.ctx_mgr.get_context(agent_id)

            # 检查模型配置，判定真实状态
            mcfg = model_configs.get(agent_id, {})
            has_api_key = bool(mcfg.get("api_key_masked", ""))
            is_enabled = mcfg.get("enabled", False)
            has_model = bool(mcfg.get("model", ""))
            is_registered = agent_id in state.orchestrator._agents

            # 状态判定: online(配置完整+启用), busy(部分配置), offline(未配置/禁用)
            if is_enabled and has_api_key and has_model and is_registered:
                real_status = "online"
            elif is_registered and (has_api_key or has_model):
                real_status = "busy"
            else:
                real_status = "offline"

            agents_list.append({
                "id": agent_id,
                "key": agent_key,
                "name": agent_def.get("name", ""),
                "persona": agent_def.get("persona", ""),
                "icon": agent_def.get("icon", ""),
                "role": agent_def.get("role", ""),
                "parallel_group": agent_def.get("parallel_group", ""),
                "has_sandbox": agent_def.get("sandbox", False),
                "sandbox_active": sandbox_path is not None,
                "context": ctx.get_progress_summary() if ctx else None,
                # 新增模型配置状态字段
                "model_status": {
                    "status": real_status,
                    "enabled": is_enabled,
                    "has_api_key": has_api_key,
                    "has_model": has_model,
                    "provider": mcfg.get("provider", ""),
                    "model": mcfg.get("model", ""),
                    "is_registered": is_registered,
                },
            })

        return ApiResponse(data=agents_list)

    @app.get("/api/agents/{agent_id}")
    async def get_agent_detail(agent_id: str):
        """获取Agent详细信息"""
        state = get_state()

        # 从规则加载器获取详细信息
        agent_spec = None
        if state.rule_loader and state.inventory:
            for aid, spec in state.inventory.agents.items():
                if agent_id in aid:
                    agent_spec = spec.to_dict()
                    break

        # 沙盒信息
        sandbox_path = state.orchestrator.sandbox_mgr.get_sandbox_path(agent_id)
        sandbox_info = None
        if sandbox_path:
            sandbox_info = {
                "path": str(sandbox_path),
                "dirs": [d.name for d in sandbox_path.iterdir() if d.is_dir()]
            }

        # 操作日志
        op_logs = state.orchestrator.sandbox_mgr.get_operation_log(agent_id=agent_id)

        # 上下文
        ctx = state.orchestrator.ctx_mgr.get_context(agent_id)

        return ApiResponse(data={
            "agent_id": agent_id,
            "spec": agent_spec,
            "sandbox": sandbox_info,
            "context": ctx.get_progress_summary() if ctx else None,
            "operation_logs": op_logs[-20:],  # 最近20条
        })

    # ── 流水线 API ───────────────────────────────────
    @app.get("/api/pipelines")
    async def list_pipelines():
        """获取所有流水线"""
        state = get_state()
        active = state.orchestrator.pipeline_engine.get_active_pipelines()

        # 也获取已完成的
        all_pipelines = []
        for pid, inst in state.orchestrator.pipeline_engine._instances.items():
            all_pipelines.append(inst.to_dict())

        return ApiResponse(data={
            "active": active,
            "all": all_pipelines,
            "total": len(all_pipelines)
        })

    @app.get("/api/pipelines/definitions")
    async def get_pipeline_definitions():
        """获取流水线定义"""
        from core.pipeline import PipelineEngine

        definitions = {}
        for req_type, defn in PipelineEngine.PIPELINE_DEFINITIONS.items():
            definitions[req_type.value] = {
                "name": defn["name"],
                "stages": [s.to_dict() for s in defn["stages"]]
            }

        quality_gates = {}
        for gate_name, gate in PipelineEngine.QUALITY_GATES.items():
            quality_gates[gate_name] = {
                "name": gate.name,
                "from_stage": gate.from_stage,
                "to_stage": gate.to_stage,
                "check_items": gate.check_items
            }

        return ApiResponse(data={
            "definitions": definitions,
            "quality_gates": quality_gates
        })

    @app.post("/api/upload")
    async def upload_files(files: List[UploadFile] = File(...)):
        """上传文件并提取文本内容，返回合并后的文件内容供 Pipeline 使用"""
        import shutil, zipfile, io

        UPLOAD_DIR = PROJECT_ROOT / "uploads"
        UPLOAD_DIR.mkdir(exist_ok=True)

        results = []
        combined_text = []

        for upload_file in files:
            try:
                content = await upload_file.read()
                filename = upload_file.filename or "unknown"
                suffix = Path(filename).suffix.lower()

                # 保存原始文件
                save_path = UPLOAD_DIR / filename
                with open(save_path, "wb") as f:
                    f.write(content)

                file_text = ""

                if suffix in (".txt", ".md"):
                    # 纯文本/Markdown 直接读取
                    file_text = content.decode("utf-8", errors="ignore")

                elif suffix == ".pdf":
                    # 尝试用 pypdf 提取文本
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(io.BytesIO(content))
                        pages = [page.extract_text() or "" for page in reader.pages]
                        file_text = "\n".join(pages)
                    except ImportError:
                        file_text = f"[PDF文件: {filename}，请安装 pypdf 以提取内容]"
                    except Exception as e:
                        file_text = f"[PDF解析失败: {e}]"

                elif suffix in (".docx", ".doc"):
                    # 尝试用 python-docx 提取文本
                    try:
                        import docx
                        doc = docx.Document(io.BytesIO(content))
                        file_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    except ImportError:
                        file_text = f"[Word文件: {filename}，请安装 python-docx 以提取内容]"
                    except Exception as e:
                        file_text = f"[Word解析失败: {e}]"

                elif suffix in (".zip", ".rar"):
                    # 解压 zip 并提取文本文件内容
                    try:
                        extracted = []
                        with zipfile.ZipFile(io.BytesIO(content)) as zf:
                            for name in zf.namelist():
                                if Path(name).suffix.lower() in (".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml", ".cs", ".cpp", ".h"):
                                    try:
                                        text = zf.read(name).decode("utf-8", errors="ignore")
                                        extracted.append(f"=== {name} ===\n{text}")
                                    except Exception:
                                        pass
                        file_text = "\n\n".join(extracted) if extracted else f"[压缩包: {filename}，未找到可读文本文件]"
                    except Exception as e:
                        file_text = f"[压缩包解析失败: {e}]"

                else:
                    # 尝试作为文本读取
                    try:
                        file_text = content.decode("utf-8", errors="ignore")
                    except Exception:
                        file_text = f"[二进制文件: {filename}]"

                results.append({
                    "filename": filename,
                    "size": len(content),
                    "text_length": len(file_text),
                    "status": "ok"
                })

                if file_text.strip():
                    combined_text.append(f"【文件: {filename}】\n{file_text.strip()}")

            except Exception as e:
                results.append({
                    "filename": upload_file.filename or "unknown",
                    "status": "error",
                    "error": str(e)
                })

        return ApiResponse(
            message=f"成功上传 {len(results)} 个文件",
            data={
                "files": results,
                "combined_text": "\n\n".join(combined_text),
                "total_text_length": sum(len(t) for t in combined_text)
            }
        )

    @app.post("/api/pipelines/create-with-files")
    async def create_pipeline_with_files(
        background_tasks: BackgroundTasks,
        user_input: str = Form(""),
        req_type: Optional[str] = Form(None),
        req_scale: Optional[str] = Form(None),
        files: List[UploadFile] = File(default=[])
    ):
        """支持文件上传的流水线创建接口（multipart/form-data）"""
        import shutil, zipfile, io

        UPLOAD_DIR = PROJECT_ROOT / "uploads"
        UPLOAD_DIR.mkdir(exist_ok=True)

        # 提取所有上传文件的文本内容
        file_texts = []
        for upload_file in files:
            try:
                content = await upload_file.read()
                filename = upload_file.filename or "unknown"
                suffix = Path(filename).suffix.lower()

                # 保存文件
                save_path = UPLOAD_DIR / filename
                with open(save_path, "wb") as f:
                    f.write(content)

                file_text = ""
                if suffix in (".txt", ".md"):
                    file_text = content.decode("utf-8", errors="ignore")
                elif suffix == ".pdf":
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(io.BytesIO(content))
                        file_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                    except Exception:
                        file_text = f"[PDF: {filename}]"
                elif suffix in (".docx", ".doc"):
                    try:
                        import docx
                        doc = docx.Document(io.BytesIO(content))
                        file_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    except Exception:
                        file_text = f"[Word: {filename}]"
                elif suffix in (".zip", ".rar"):
                    try:
                        extracted = []
                        with zipfile.ZipFile(io.BytesIO(content)) as zf:
                            for name in zf.namelist():
                                if Path(name).suffix.lower() in (".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml", ".cs", ".cpp", ".h"):
                                    try:
                                        text = zf.read(name).decode("utf-8", errors="ignore")
                                        extracted.append(f"=== {name} ===\n{text}")
                                    except Exception:
                                        pass
                        file_text = "\n\n".join(extracted)
                    except Exception:
                        file_text = f"[ZIP: {filename}]"
                else:
                    try:
                        file_text = content.decode("utf-8", errors="ignore")
                    except Exception:
                        pass

                if file_text.strip():
                    file_texts.append(f"【文件: {filename}】\n{file_text.strip()}")

            except Exception as e:
                logger.warning(f"处理上传文件失败: {upload_file.filename}: {e}")

        # 合并用户输入和文件内容作为需求描述
        parts = []
        if user_input.strip():
            parts.append(user_input.strip())
        if file_texts:
            parts.append("\n\n".join(file_texts))

        combined_input = "\n\n---\n\n".join(parts) if parts else "（上传文件，无文字描述）"

        # 创建 Pipeline
        state = get_state()
        from core.pipeline import ReqType, ReqScale

        try:
            _req_type = ReqType(req_type or "FEATURE")
            _req_scale = ReqScale(req_scale or "M")
        except (ValueError, KeyError):
            raise HTTPException(status_code=400, detail="无效的需求类型或规模")

        pipeline = state.orchestrator.pipeline_engine.create_pipeline(
            req_id=f"REQ-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            req_type=_req_type,
            req_scale=_req_scale,
            req_name=user_input.strip() or (file_texts[0][:50] + "..." if file_texts else "文件上传需求")
        )

        pipeline.status = "running"
        pipeline.started_at = datetime.now().isoformat()
        state.orchestrator._active_pipelines[pipeline.pipeline_id] = pipeline

        # 持久化到数据库
        try:
            from core.database import save_pipeline, save_uploaded_file, log_event
            save_pipeline(pipeline.to_dict(), user_input=combined_input)
            # 记录每个上传文件
            for upload_file in files:
                await upload_file.seek(0)
                raw = await upload_file.read()
                save_uploaded_file(
                    pipeline_id=pipeline.pipeline_id,
                    filename=upload_file.filename or "unknown",
                    original_path=str(UPLOAD_DIR / (upload_file.filename or "unknown")),
                    file_size=len(raw),
                    mime_type=upload_file.content_type,
                    extracted_text="",  # 已合并到 combined_input
                )
            log_event("pipeline_created_with_files",
                      f"流水线 {pipeline.pipeline_id} 已创建（含 {len(files)} 个文件）",
                      pipeline_id=pipeline.pipeline_id)
        except Exception as _db_err:
            logger.warning(f"数据库写入失败: {_db_err}")

        state.orchestrator._record_log(
            "pipeline_created",
            f"流水线 {pipeline.pipeline_id} 已创建（含 {len(file_texts)} 个文件）: {combined_input[:50]}",
            pipeline_id=pipeline.pipeline_id
        )

        background_tasks.add_task(
            state.execute_pipeline_async,
            pipeline.pipeline_id,
            combined_input
        )

        return ApiResponse(
            message=f"流水线 {pipeline.pipeline_id} 已创建并开始执行",
            data={
                **pipeline.to_dict(),
                "uploaded_files": len(file_texts)
            }
        )

    @app.post("/api/pipelines/create")
    async def create_pipeline(req: RequirementRequest, background_tasks: BackgroundTasks):
        """创建并异步执行流水线"""
        state = get_state()
        from core.pipeline import ReqType, ReqScale

        try:
            req_type = ReqType(req.req_type or "FEATURE")
            req_scale = ReqScale(req.req_scale or "M")
        except (ValueError, KeyError):
            raise HTTPException(status_code=400, detail="无效的需求类型或规模")

        # 创建流水线
        pipeline = state.orchestrator.pipeline_engine.create_pipeline(
            req_id=f"REQ-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            req_type=req_type,
            req_scale=req_scale,
            req_name=req.user_input
        )

        # 立即设置为 running 状态（让前端看到进度）
        pipeline.status = "running"
        pipeline.started_at = datetime.now().isoformat()
        state.orchestrator._active_pipelines[pipeline.pipeline_id] = pipeline

        # 持久化到数据库
        try:
            from core.database import save_pipeline, log_event
            save_pipeline(pipeline.to_dict(), user_input=req.user_input)
            log_event("pipeline_created",
                      f"流水线 {pipeline.pipeline_id} 已创建: {req.user_input[:50]}",
                      pipeline_id=pipeline.pipeline_id)
        except Exception as _db_err:
            logger.warning(f"数据库写入失败: {_db_err}")

        # 记录日志
        state.orchestrator._record_log(
            "pipeline_created",
            f"流水线 {pipeline.pipeline_id} 已创建并开始执行: {req.user_input[:50]}",
            pipeline_id=pipeline.pipeline_id
        )

        # 异步执行流水线（不阻塞 API 返回）
        background_tasks.add_task(
            state.execute_pipeline_async,
            pipeline.pipeline_id,
            req.user_input
        )

        return ApiResponse(
            message=f"流水线 {pipeline.pipeline_id} 已创建并开始执行",
            data=pipeline.to_dict()
        )

    # ── 流水线删除 API ───────────────────────────────
    @app.delete("/api/pipelines/{pipeline_id}")
    async def delete_pipeline_endpoint(pipeline_id: str):
        """删除流水线（彻底清除：引擎实例 + 活跃列表 + 异步任务记录）"""
        state = get_state()

        # 1. 从流水线引擎中删除
        deleted = state.orchestrator.pipeline_engine.delete_pipeline(pipeline_id)

        # 2. 从活跃流水线字典中移除
        state.orchestrator._active_pipelines.pop(pipeline_id, None)

        # 3. 从异步执行记录中移除
        state._running_pipelines.pop(pipeline_id, None)

        # 4. 清除该项目的所有活动日志
        state.orchestrator.delete_pipeline_logs(pipeline_id)

        if not deleted:
            raise HTTPException(status_code=404, detail=f"流水线 {pipeline_id} 不存在")

        # 数据库同步删除
        try:
            from core.database import delete_pipeline_db
            delete_pipeline_db(pipeline_id)
        except Exception as _db_err:
            logger.warning(f"数据库删除失败: {_db_err}")

        state.orchestrator._record_log(
            "pipeline_deleted",
            f"流水线 {pipeline_id} 已被用户删除",
            pipeline_id=pipeline_id
        )

        return ApiResponse(message=f"流水线 {pipeline_id} 已删除")

    # ── 流水线重命名 API ─────────────────────────────
    @app.put("/api/pipelines/{pipeline_id}/rename")
    async def rename_pipeline_endpoint(pipeline_id: str, req: RenameRequest):
        """重命名流水线"""
        state = get_state()

        if not req.name or not req.name.strip():
            raise HTTPException(status_code=400, detail="名称不能为空")

        instance = state.orchestrator.pipeline_engine.rename_pipeline(
            pipeline_id, req.name.strip()
        )

        if not instance:
            raise HTTPException(status_code=404, detail=f"流水线 {pipeline_id} 不存在")

        state.orchestrator._record_log(
            "pipeline_renamed",
            f"流水线 {pipeline_id} 已重命名为: {req.name.strip()}",
            pipeline_id=pipeline_id
        )

        # 数据库同步重命名
        try:
            from core.database import rename_pipeline_db
            rename_pipeline_db(pipeline_id, req.name.strip())
        except Exception as _db_err:
            logger.warning(f"数据库重命名失败: {_db_err}")

        return ApiResponse(
            message=f"流水线 {pipeline_id} 已重命名",
            data=instance.to_dict()
        )

    # ── 流水线日志 API（按项目）──────────────────────
    @app.get("/api/pipelines/{pipeline_id}/logs")
    async def get_pipeline_logs(pipeline_id: str, count: int = Query(50, ge=1, le=500)):
        """获取指定流水线的活动日志"""
        state = get_state()

        # 验证流水线存在
        pipeline = state.orchestrator.pipeline_engine.get_pipeline(pipeline_id)
        if not pipeline:
            raise HTTPException(status_code=404, detail=f"流水线 {pipeline_id} 不存在")

        logs = state.orchestrator.get_pipeline_logs(pipeline_id, count)
        return ApiResponse(data=logs)

    # ── 消息队列 API ─────────────────────────────────
    @app.get("/api/messages/stats")
    async def get_message_stats():
        """获取消息队列统计"""
        state = get_state()
        stats = state.orchestrator.mq.get_queue_stats()
        return ApiResponse(data=stats)

    @app.get("/api/messages/{agent_id}")
    async def get_agent_messages(agent_id: str, limit: int = Query(20, ge=1, le=100)):
        """获取Agent的消息"""
        state = get_state()
        messages = state.orchestrator.mq.receive(agent_id, limit=limit)
        return ApiResponse(data=[m.to_dict() for m in messages])

    @app.post("/api/messages/send")
    async def send_message(msg: MessageRequest):
        """发送消息"""
        state = get_state()
        from core.message_queue import Message, MessageType, MessagePriority

        message = Message(
            from_agent=msg.from_agent,
            to_agent=msg.to_agent,
            msg_type=msg.msg_type,
            priority=MessagePriority.NORMAL.value,
            payload=msg.payload
        )

        msg_id = state.orchestrator.mq.send(message)
        return ApiResponse(
            message=f"消息已发送: {msg_id}",
            data={"msg_id": msg_id}
        )

    # ── 沙盒管理 API ─────────────────────────────────
    @app.get("/api/sandboxes")
    async def list_sandboxes():
        """获取所有沙盒"""
        state = get_state()
        sandboxes = state.orchestrator.sandbox_mgr.list_active_sandboxes()
        return ApiResponse(data=sandboxes)

    @app.get("/api/sandboxes/{agent_id}/logs")
    async def get_sandbox_logs(agent_id: str):
        """获取沙盒操作日志"""
        state = get_state()
        logs = state.orchestrator.sandbox_mgr.get_operation_log(agent_id=agent_id)
        return ApiResponse(data=logs)

    @app.post("/api/sandboxes/check-access")
    async def check_access(agent_id: str = Query(...),
                           file_path: str = Query(...),
                           operation: str = Query(...)):
        """检查文件访问权限"""
        state = get_state()
        allowed, reason = state.orchestrator.sandbox_mgr.check_access(
            agent_id, file_path, operation
        )
        return ApiResponse(data={
            "allowed": allowed,
            "reason": reason,
            "agent_id": agent_id,
            "file_path": file_path,
            "operation": operation
        })

    # ── 规则资产 API ─────────────────────────────────
    @app.get("/api/inventory")
    async def get_inventory():
        """获取规则资产清单"""
        state = get_state()
        if state.inventory:
            return ApiResponse(data=state.inventory.to_dict())
        return ApiResponse(success=False, message="未加载规则资产")

    @app.post("/api/inventory/rescan")
    async def rescan_inventory():
        """重新扫描规则资产"""
        state = get_state()
        state.inventory = state.rule_loader.scan_all()
        issues = state.rule_loader.validate_inventory()
        return ApiResponse(
            message="重新扫描完成",
            data={
                "statistics": state.inventory.statistics,
                "issues": issues
            }
        )

    # ── 智能体规则内容 API ────────────────────────────
    @app.get("/api/agents/{agent_id}/rules")
    async def get_agent_rules(agent_id: str):
        """获取智能体的完整规则内容（入口文件、步骤、模板）"""
        state = get_state()
        if not state.rule_loader or not state.inventory:
            return ApiResponse(success=False, message="规则加载器未初始化")

        # 在 inventory 中查找匹配的 agent
        # agent_id 来自 system.yaml (如 "00_producer"), inventory key 来自 rules/ (如 "00_制作人Agent")
        # 匹配策略: 1) 精确匹配 2) 数字前缀匹配 (提取前缀数字部分做比对)
        import re as _re
        agent_spec = None
        matched_id = None
        agent_num_match = _re.match(r'^(\d+)', agent_id)
        agent_num_prefix = agent_num_match.group(1) if agent_num_match else None
        for aid, spec in state.inventory.agents.items():
            if aid == agent_id:
                agent_spec = spec
                matched_id = aid
                break
            aid_num_match = _re.match(r'^(\d+)', aid)
            if aid_num_match and agent_num_prefix and aid_num_match.group(1) == agent_num_prefix:
                agent_spec = spec
                matched_id = aid
                break

        if not agent_spec:
            raise HTTPException(status_code=404, detail=f"Agent {agent_id} 规则不存在")

        # 读取入口文件内容
        entry_content = state.rule_loader.load_agent_entry_content(matched_id) or ""

        # 读取所有步骤文件
        steps = []
        for step_path in agent_spec.steps:
            content = state.rule_loader.load_step_content(step_path) or ""
            # 从路径提取步骤名
            step_name = Path(step_path).stem
            steps.append({
                "path": step_path,
                "name": step_name,
                "content": content,
            })

        # 读取所有模板文件
        templates = []
        for tmpl_path in agent_spec.templates:
            content = state.rule_loader.load_template_content(tmpl_path) or ""
            tmpl_name = Path(tmpl_path).stem
            templates.append({
                "path": tmpl_path,
                "name": tmpl_name,
                "content": content,
            })

        return ApiResponse(data={
            "agent_id": matched_id,
            "agent_name": agent_spec.agent_name,
            "title": agent_spec.title,
            "entry_content": entry_content,
            "steps": steps,
            "templates": templates,
            "file_count": agent_spec.file_count,
            "directory": agent_spec.directory,
        })

    # ── 团队能力详情 API ──────────────────────────────
    @app.get("/api/capabilities")
    async def get_capabilities():
        """获取团队特色能力的详细信息（从真实配置派生）"""
        state = get_state()
        from utils.file_ops import safe_yaml_read

        config = safe_yaml_read(state.project_root / "config" / "system.yaml") or {}
        agents_config = config.get("agents", {})
        parallel_groups = config.get("parallel_groups", {})

        # 统计沙盒
        sandbox_agents = [
            {"id": v.get("id", k), "name": v.get("name", k)}
            for k, v in agents_config.items() if v.get("sandbox", False)
        ]

        # 质量门禁
        quality_gates = []
        if state.inventory:
            # 从 rules/ 中读取门禁规则
            guard_content = state.rule_loader.load_step_content("rule_guards.md")
            if guard_content:
                quality_gates.append({
                    "name": "质量门禁体系",
                    "description": guard_content[:500],
                })

        # 并行组
        parallel_info = []
        for gname, gdef in parallel_groups.items():
            parallel_info.append({
                "group": gname,
                "agents": gdef.get("agents", []),
                "max_concurrent": gdef.get("max_concurrent", 1),
            })

        # 构建能力详情
        capabilities = [
            {
                "id": "quality_gates",
                "icon": "🛡️",
                "title": "3 道质量门禁",
                "description": "策划门禁 (21项)、技术门禁 (18项)、测试门禁 (12项)",
                "detail": "通过三层质量检查确保每个阶段产出物符合标准，任何一道门禁不通过都会触发返工流程。",
                "data": quality_gates,
            },
            {
                "id": "bug_fix",
                "icon": "🐛",
                "title": "自动 Bug 修复循环",
                "description": "QA 发现 Bug 自动回滚给程序，最多 3 轮修复循环",
                "detail": "QA智能体发现缺陷后自动创建Bug报告，系统回滚至程序员Agent，最多循环3轮。超出后升级至主程Agent介入。",
                "data": {"max_rounds": 3, "escalation": "03_tech_lead"},
            },
            {
                "id": "parallel",
                "icon": "⚡",
                "title": "并行协作架构",
                "description": f"{len(parallel_groups)} 个并行组同时工作，大幅提升开发效率",
                "detail": "基于独立沙盒和消息队列，多个Agent可同时执行不同阶段的任务，通过并行组配置管理并发度。",
                "data": parallel_info,
            },
            {
                "id": "sandbox",
                "icon": "📦",
                "title": "沙盒隔离环境",
                "description": f"{len(sandbox_agents)} 位成员拥有独立工作空间，互不干扰",
                "detail": "每个Agent拥有独立的文件系统沙盒，通过权限矩阵严格控制读写范围。沙盒之间完全隔离，只能通过消息队列通信。",
                "data": sandbox_agents,
            },
            {
                "id": "review",
                "icon": "🔍",
                "title": "对抗审查机制",
                "description": "程序员自我审查 + 主程交叉审查，双重保障",
                "detail": "程序员Agent编写代码后先进行自我审查（包括SOLID原则、性能、安全），再由主程Agent进行交叉审查，确保代码质量。",
                "data": {"self_review": "04_programmer", "cross_review": "03_tech_lead"},
            },
            {
                "id": "deliverables",
                "icon": "📄",
                "title": "完整交付物",
                "description": "策划案、技术设计、代码、测试报告一应俱全",
                "detail": "每个阶段都会产出标准化文档：需求分析报告、策划案、技术设计文档、源代码、单元测试、测试报告。所有交付物可追溯。",
                "data": {
                    "stages": ["需求分析", "策划案", "技术设计", "源代码", "测试用例", "测试报告"]
                },
            },
        ]

        return ApiResponse(data=capabilities)

    # ── 自定义规则模板 CRUD API ────────────────────────
    @app.get("/api/agents/{agent_id}/custom-rules")
    async def list_custom_rules(agent_id: str):
        """列出智能体的自定义规则模板"""
        state = get_state()
        custom_dir = state.project_root / "rules" / "agents" / "custom" / agent_id
        if not custom_dir.exists():
            return ApiResponse(data=[])

        rules = []
        for f in sorted(custom_dir.glob("*.md")):
            content = f.read_text(encoding="utf-8")
            from utils.file_ops import extract_title
            title = extract_title(content)
            rules.append({
                "filename": f.name,
                "title": title or f.stem,
                "content": content,
                "size": f.stat().st_size,
                "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
            })
        return ApiResponse(data=rules)

    @app.post("/api/agents/{agent_id}/custom-rules")
    async def create_custom_rule(agent_id: str, req: CustomRuleRequest):
        """为智能体创建自定义规则模板"""
        state = get_state()
        custom_dir = state.project_root / "rules" / "agents" / "custom" / agent_id
        custom_dir.mkdir(parents=True, exist_ok=True)

        filename = req.filename or f"custom_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        if not filename.endswith(".md"):
            filename += ".md"

        filepath = custom_dir / filename
        filepath.write_text(req.content, encoding="utf-8")

        return ApiResponse(
            message=f"规则模板 {filename} 已创建",
            data={"filename": filename, "path": str(filepath.relative_to(state.project_root))}
        )

    @app.put("/api/agents/{agent_id}/custom-rules/{filename}")
    async def update_custom_rule(agent_id: str, filename: str, req: CustomRuleRequest):
        """更新自定义规则模板"""
        state = get_state()
        filepath = state.project_root / "rules" / "agents" / "custom" / agent_id / filename
        if not filepath.exists():
            raise HTTPException(status_code=404, detail="规则文件不存在")

        filepath.write_text(req.content, encoding="utf-8")
        return ApiResponse(message=f"规则模板 {filename} 已更新")

    @app.delete("/api/agents/{agent_id}/custom-rules/{filename}")
    async def delete_custom_rule(agent_id: str, filename: str):
        """删除自定义规则模板"""
        state = get_state()
        filepath = state.project_root / "rules" / "agents" / "custom" / agent_id / filename
        if not filepath.exists():
            raise HTTPException(status_code=404, detail="规则文件不存在")

        filepath.unlink()
        return ApiResponse(message=f"规则模板 {filename} 已删除")

    # ── 新智能体上传/创建 API ─────────────────────────
    @app.post("/api/agents/create")
    async def create_new_agent(req: NewAgentRequest):
        """创建新的智能体（规则文件 + 目录结构 + 沙盒 + 注册）"""
        state = get_state()

        agents_dir = state.project_root / "rules" / "agents"
        if not agents_dir.exists():
            raise HTTPException(status_code=500, detail="rules/agents 目录不存在")

        # 计算下一个编号
        existing_nums = []
        import re as _re
        for f in agents_dir.glob("*_*Agent.md"):
            m = _re.match(r'^(\d{2})_', f.stem)
            if m:
                existing_nums.append(int(m.group(1)))
        next_num = max(existing_nums, default=7) + 1
        agent_id = f"{next_num:02d}_{req.agent_name}Agent"

        # 创建入口文件
        entry_content = req.entry_content or f"""---
type: manual
---
# {req.agent_name} Agent 职能规范

> **版本**: v1.0
> **创建时间**: {datetime.now().strftime('%Y-%m-%d')}
> **角色**: {req.role or '自定义智能体'}

---

## 🎯 角色定位

{req.persona or f'你是一名专业的{req.role or "AI助手"}，负责完成用户分配的任务。'}

### 🎭 人格档案

| 属性 | 值 |
|------|-----|
| **名称** | {req.agent_icon} {req.agent_name} |
| **图标** | {req.agent_icon} |
| **角色** | {req.role} |

### 🔐 工具权限声明

| 权限类型 | 允许范围 | 禁止范围 |
|---------|---------|---------|
| **读取** | 规则文件 | 其他Agent工作区 |
| **修改** | 自身工作区文件 | 系统配置文件 |

---

## 📋 核心职责

*请在此定义该智能体的核心工作职责*

---

## 🔄 步骤文件索引

*暂无步骤文件，可在下方目录中添加 step-XX_*.md 文件*
"""

        entry_path = agents_dir / f"{agent_id}.md"
        entry_path.write_text(entry_content, encoding="utf-8")

        # 创建子目录和模板目录
        agent_dir = agents_dir / agent_id
        agent_dir.mkdir(exist_ok=True)
        (agent_dir / "templates").mkdir(exist_ok=True)

        # 创建沙盒环境
        sandbox_created = False
        if req.sandbox:
            try:
                from core.sandbox import SandboxConfig
                from utils.file_ops import safe_yaml_read
                config = safe_yaml_read(state.project_root / "config" / "system.yaml") or {}
                parallel_config = config.get("parallel_architecture", {})
                sc = SandboxConfig(
                    agent_id=agent_id,
                    agent_name=req.agent_name,
                    sandbox_root=parallel_config.get("sandbox_root", ".sandboxes"),
                    read_permissions=["rules/**"],
                    write_permissions=[".GameDev/**"]
                )
                state.orchestrator.sandbox_mgr.create_sandbox(sc)
                sandbox_created = True
            except Exception as e:
                logger.warning(f"为新Agent {agent_id} 创建沙盒失败: {e}")

        # 初始化智能体配置数据目录
        agent_data_dir = state.project_root / "config" / "agents" / agent_id
        agent_data_dir.mkdir(parents=True, exist_ok=True)
        # 创建默认配置文件
        agent_config_file = agent_data_dir / "config.json"
        if not agent_config_file.exists():
            default_cfg = {
                "agent_id": agent_id,
                "agent_name": req.agent_name,
                "icon": req.agent_icon,
                "role": req.role,
                "group": req.group,
                "plugins": [],
                "mcp_servers": [],
                "integrations": [],
                "skills": [],
                "memory": [],
                "created_at": datetime.now().isoformat(),
            }
            agent_config_file.write_text(json.dumps(default_cfg, ensure_ascii=False, indent=2), encoding="utf-8")

        # 重新扫描规则资产
        state.inventory = state.rule_loader.scan_all()

        return ApiResponse(
            message=f"智能体 {req.agent_name} 创建成功",
            data={
                "agent_id": agent_id,
                "entry_file": str(entry_path.relative_to(state.project_root)),
                "directory": str(agent_dir.relative_to(state.project_root)),
                "sandbox_created": sandbox_created,
            }
        )

    # ── 智能体完整配置面板 API (仿 CodeBuddy) ─────────

    def _get_agent_config_path(agent_id: str):
        """获取智能体配置文件路径"""
        state = get_state()
        cfg_path = state.project_root / "config" / "agents" / agent_id / "config.json"
        return cfg_path

    def _load_agent_panel_config(agent_id: str) -> dict:
        """读取智能体面板配置"""
        cfg_path = _get_agent_config_path(agent_id)
        if cfg_path.exists():
            try:
                return json.loads(cfg_path.read_text(encoding="utf-8"))
            except Exception:
                pass
        # 返回默认配置
        return {
            "agent_id": agent_id,
            "plugins": [],
            "mcp_servers": [],
            "integrations": [],
            "skills": [],
            "memory": [],
        }

    def _save_agent_panel_config(agent_id: str, config: dict):
        """保存智能体面板配置"""
        cfg_path = _get_agent_config_path(agent_id)
        cfg_path.parent.mkdir(parents=True, exist_ok=True)
        cfg_path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    @app.get("/api/agents/{agent_id}/panel")
    async def get_agent_panel(agent_id: str):
        """获取智能体完整面板配置（所有 Tab 数据）"""
        state = get_state()
        panel = _load_agent_panel_config(agent_id)

        # 加载规则数据 (rules tab)
        # 使用数字前缀匹配: system.yaml id (如 "00_producer") → rules/ id (如 "00_制作人Agent")
        import re as _re
        rules_data = {"entry_content": "", "steps": [], "templates": [], "agent_id": agent_id, "agent_name": "", "title": "", "file_count": 0, "directory": ""}
        if state.rule_loader and state.inventory:
            agent_num_match = _re.match(r'^(\d+)', agent_id)
            agent_num_prefix = agent_num_match.group(1) if agent_num_match else None
            for aid, spec in state.inventory.agents.items():
                aid_num_match = _re.match(r'^(\d+)', aid)
                matched = (aid == agent_id) or (aid_num_match and agent_num_prefix and aid_num_match.group(1) == agent_num_prefix)
                if matched:
                    entry = state.rule_loader.load_agent_entry_content(aid) or ""
                    steps = []
                    for sp in spec.steps:
                        c = state.rule_loader.load_step_content(sp) or ""
                        steps.append({"path": sp, "name": Path(sp).stem, "content": c})
                    templates = []
                    for tp in spec.templates:
                        c = state.rule_loader.load_template_content(tp) or ""
                        templates.append({"path": tp, "name": Path(tp).stem, "content": c})
                    rules_data = {
                        "agent_id": aid,
                        "agent_name": spec.agent_name,
                        "title": spec.title,
                        "entry_content": entry,
                        "steps": steps,
                        "templates": templates,
                        "file_count": spec.file_count,
                        "directory": spec.directory,
                    }
                    break

        # 加载自定义规则
        custom_rules = []
        custom_dir = state.project_root / "rules" / "agents" / "custom" / agent_id
        if custom_dir.exists():
            from utils.file_ops import extract_title
            for f in sorted(custom_dir.glob("*.md")):
                content = f.read_text(encoding="utf-8")
                title = extract_title(content)
                custom_rules.append({
                    "filename": f.name,
                    "title": title or f.stem,
                    "content": content,
                    "size": f.stat().st_size,
                    "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
                })

        # 加载可用技能（从 inventory）
        available_skills = []
        if state.inventory:
            for sid, sspec in state.inventory.skills.items():
                bound = any(s.get("skill_id") == sid for s in panel.get("skills", []))
                available_skills.append({
                    **sspec.to_dict(),
                    "bound": bound,
                    "enabled": bound,
                })

        # 加载可用插件（系统内置）
        available_plugins = _get_system_plugins()

        # 加载可用集成
        available_integrations = _get_system_integrations()

        return ApiResponse(data={
            "config": panel,
            "rules": rules_data,
            "custom_rules": custom_rules,
            "available_skills": available_skills,
            "available_plugins": available_plugins,
            "available_integrations": available_integrations,
        })

    def _get_system_plugins() -> list:
        """系统内置插件列表"""
        return [
            {"id": "code_review", "name": "代码审查", "icon": "🔍", "desc": "自动代码审查，检测常见问题和最佳实践违规", "category": "quality", "tags": ["rules"]},
            {"id": "auto_test", "name": "自动测试", "icon": "🧪", "desc": "根据代码变更自动生成和执行单元测试", "category": "testing", "tags": ["skills"]},
            {"id": "doc_gen", "name": "文档生成", "icon": "📄", "desc": "自动生成 API 文档、技术设计文档和 README", "category": "productivity", "tags": ["skills"]},
            {"id": "perf_monitor", "name": "性能监控", "icon": "📊", "desc": "实时监控代码性能，检测内存泄漏和性能瓶颈", "category": "monitoring", "tags": ["skills", "rules"]},
            {"id": "security_scan", "name": "安全扫描", "icon": "🛡️", "desc": "扫描代码安全漏洞，支持 OWASP Top 10 检测", "category": "security", "tags": ["rules"]},
            {"id": "git_flow", "name": "Git工作流", "icon": "🔀", "desc": "自动化 Git 操作：分支管理、commit 规范、PR 创建", "category": "workflow", "tags": ["skills"]},
            {"id": "dependency_mgr", "name": "依赖管理", "icon": "📦", "desc": "自动检查和更新项目依赖，检测安全漏洞", "category": "maintenance", "tags": ["skills"]},
            {"id": "i18n", "name": "国际化", "icon": "🌍", "desc": "自动提取和管理多语言文本，支持翻译工作流", "category": "productivity", "tags": ["skills"]},
        ]

    def _get_system_integrations() -> list:
        """系统内置集成列表"""
        return [
            {"id": "github", "name": "GitHub", "icon": "🐙", "desc": "连接 GitHub 仓库，同步 issues、PR 和代码", "status": "available"},
            {"id": "jira", "name": "Jira", "icon": "📋", "desc": "同步 Jira 任务和需求跟踪", "status": "available"},
            {"id": "slack", "name": "Slack", "icon": "💬", "desc": "将团队通知和进度更新推送到 Slack 频道", "status": "available"},
            {"id": "unity_cloud", "name": "Unity Cloud", "icon": "🎮", "desc": "连接 Unity Cloud Build，自动构建和部署", "status": "available"},
            {"id": "docker", "name": "Docker", "icon": "🐳", "desc": "容器化构建和部署环境管理", "status": "available"},
            {"id": "ci_cd", "name": "CI/CD Pipeline", "icon": "🔄", "desc": "集成持续集成/持续部署工具链", "status": "available"},
        ]

    # ── 插件 Tab API ──────────────────────────────────
    @app.get("/api/agents/{agent_id}/plugins")
    async def get_agent_plugins(agent_id: str):
        """获取智能体已启用的插件列表"""
        panel = _load_agent_panel_config(agent_id)
        all_plugins = _get_system_plugins()
        enabled_ids = {p["plugin_id"] for p in panel.get("plugins", [])}
        for p in all_plugins:
            p["enabled"] = p["id"] in enabled_ids
            # 合并自定义配置
            match = next((ep for ep in panel.get("plugins", []) if ep.get("plugin_id") == p["id"]), None)
            if match:
                p["config"] = match.get("config", {})
        return ApiResponse(data=all_plugins)

    @app.put("/api/agents/{agent_id}/plugins")
    async def update_agent_plugin(agent_id: str, req: AgentPluginRequest):
        """启用/禁用/配置智能体插件"""
        panel = _load_agent_panel_config(agent_id)
        plugins = panel.get("plugins", [])
        existing = next((p for p in plugins if p.get("plugin_id") == req.plugin_id), None)
        if req.enabled:
            if existing:
                existing["config"] = req.config
            else:
                plugins.append({"plugin_id": req.plugin_id, "enabled": True, "config": req.config})
        else:
            plugins = [p for p in plugins if p.get("plugin_id") != req.plugin_id]
        panel["plugins"] = plugins
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message=f"插件 {req.plugin_id} 已{'启用' if req.enabled else '禁用'}")

    # ── MCP Tab API ───────────────────────────────────
    @app.get("/api/agents/{agent_id}/mcp")
    async def get_agent_mcp(agent_id: str):
        """获取智能体的 MCP 服务器列表"""
        panel = _load_agent_panel_config(agent_id)
        return ApiResponse(data=panel.get("mcp_servers", []))

    @app.post("/api/agents/{agent_id}/mcp")
    async def add_agent_mcp(agent_id: str, req: AgentMCPRequest):
        """添加 MCP 服务器到智能体"""
        panel = _load_agent_panel_config(agent_id)
        servers = panel.get("mcp_servers", [])
        servers.append({
            "server_name": req.server_name,
            "server_url": req.server_url,
            "enabled": req.enabled,
            "tools": req.tools,
            "added_at": datetime.now().isoformat(),
        })
        panel["mcp_servers"] = servers
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message=f"MCP 服务器 {req.server_name} 已添加")

    @app.delete("/api/agents/{agent_id}/mcp/{server_name}")
    async def remove_agent_mcp(agent_id: str, server_name: str):
        """移除 MCP 服务器"""
        panel = _load_agent_panel_config(agent_id)
        panel["mcp_servers"] = [s for s in panel.get("mcp_servers", []) if s.get("server_name") != server_name]
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message=f"MCP 服务器 {server_name} 已移除")

    # ── 技能 Tab API ──────────────────────────────────
    @app.get("/api/agents/{agent_id}/skills")
    async def get_agent_skills(agent_id: str):
        """获取智能体已绑定的技能"""
        state = get_state()
        panel = _load_agent_panel_config(agent_id)
        bound_ids = {s.get("skill_id") for s in panel.get("skills", [])}

        all_skills = []
        if state.inventory:
            for sid, sspec in state.inventory.skills.items():
                content = state.rule_loader.load_skill_content(sid) or ""
                all_skills.append({
                    **sspec.to_dict(),
                    "bound": sid in bound_ids,
                    "content_preview": content[:300] if content else "",
                })
        return ApiResponse(data=all_skills)

    @app.put("/api/agents/{agent_id}/skills")
    async def bind_agent_skill(agent_id: str, req: AgentSkillBindRequest):
        """绑定/解绑技能到智能体"""
        panel = _load_agent_panel_config(agent_id)
        skills = panel.get("skills", [])
        if req.enabled:
            if not any(s.get("skill_id") == req.skill_id for s in skills):
                skills.append({"skill_id": req.skill_id, "enabled": True, "bound_at": datetime.now().isoformat()})
        else:
            skills = [s for s in skills if s.get("skill_id") != req.skill_id]
        panel["skills"] = skills
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message=f"技能 {req.skill_id} 已{'绑定' if req.enabled else '解绑'}")

    # ── 集成 Tab API ──────────────────────────────────
    @app.get("/api/agents/{agent_id}/integrations")
    async def get_agent_integrations(agent_id: str):
        """获取智能体的集成列表"""
        panel = _load_agent_panel_config(agent_id)
        enabled_ids = {i.get("integration_id") for i in panel.get("integrations", [])}
        all_ints = _get_system_integrations()
        for it in all_ints:
            it["enabled"] = it["id"] in enabled_ids
            match = next((ei for ei in panel.get("integrations", []) if ei.get("integration_id") == it["id"]), None)
            if match:
                it["config"] = match.get("config", {})
        return ApiResponse(data=all_ints)

    @app.put("/api/agents/{agent_id}/integrations")
    async def update_agent_integration(agent_id: str, req: AgentIntegrationRequest):
        """启用/禁用集成"""
        panel = _load_agent_panel_config(agent_id)
        integrations = panel.get("integrations", [])
        if req.enabled:
            existing = next((i for i in integrations if i.get("integration_id") == req.integration_id), None)
            if existing:
                existing["config"] = req.config
            else:
                integrations.append({"integration_id": req.integration_id, "enabled": True, "config": req.config})
        else:
            integrations = [i for i in integrations if i.get("integration_id") != req.integration_id]
        panel["integrations"] = integrations
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message=f"集成 {req.integration_id} 已{'启用' if req.enabled else '禁用'}")

    # ── 记忆 Tab API ──────────────────────────────────
    @app.get("/api/agents/{agent_id}/memory")
    async def get_agent_memory(agent_id: str):
        """获取智能体记忆列表"""
        panel = _load_agent_panel_config(agent_id)
        return ApiResponse(data=panel.get("memory", []))

    @app.post("/api/agents/{agent_id}/memory")
    async def add_agent_memory(agent_id: str, req: AgentMemoryRequest):
        """添加记忆条目"""
        panel = _load_agent_panel_config(agent_id)
        memories = panel.get("memory", [])
        import uuid
        memories.append({
            "id": str(uuid.uuid4())[:8],
            "title": req.title,
            "content": req.content,
            "memory_type": req.memory_type,
            "created_at": datetime.now().isoformat(),
        })
        panel["memory"] = memories
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message="记忆已添加")

    @app.delete("/api/agents/{agent_id}/memory/{memory_id}")
    async def delete_agent_memory(agent_id: str, memory_id: str):
        """删除记忆条目"""
        panel = _load_agent_panel_config(agent_id)
        panel["memory"] = [m for m in panel.get("memory", []) if m.get("id") != memory_id]
        _save_agent_panel_config(agent_id, panel)
        return ApiResponse(message="记忆已删除")

    # ── CodeBuddy 适配器 API ─────────────────────────
    @app.get("/api/codebuddy/team")
    async def get_team_status():
        """获取CodeBuddy团队状态"""
        state = get_state()
        return ApiResponse(data=state.codebuddy_adapter.get_team_status())

    # ── 并行组 API ───────────────────────────────────
    @app.get("/api/parallel-groups")
    async def get_parallel_groups():
        """获取并行组定义"""
        from utils.file_ops import safe_yaml_read
        state = get_state()
        config = safe_yaml_read(state.project_root / "config" / "system.yaml") or {}
        groups = config.get("parallel_groups", {})
        return ApiResponse(data=groups)

    # ── 系统日志 API ─────────────────────────────────
    @app.get("/api/logs")
    async def get_system_logs(count: int = Query(50, ge=1, le=500)):
        """获取系统日志"""
        state = get_state()
        logs = state.orchestrator._log[-count:]
        return ApiResponse(data=logs)

    # ── 智能体模型配置 API ───────────────────────────
    @app.get("/api/agent-configs")
    async def get_agent_configs():
        """获取所有智能体的模型配置（脱敏）"""
        state = get_state()
        configs = state.model_config_mgr.get_all_configs_masked()
        return ApiResponse(data=configs)

    @app.get("/api/agent-configs/{agent_id}")
    async def get_agent_config(agent_id: str):
        """获取指定智能体的模型配置"""
        state = get_state()
        cfg = state.model_config_mgr.get_config(agent_id)
        if not cfg:
            raise HTTPException(status_code=404, detail=f"Agent {agent_id} 配置不存在")
        return ApiResponse(data=cfg.to_dict())

    @app.put("/api/agent-configs/{agent_id}")
    async def update_agent_config(agent_id: str, req: AgentConfigRequest):
        """更新指定智能体的模型配置"""
        state = get_state()
        updates = {k: v for k, v in req.dict().items() if v is not None}
        if not updates:
            raise HTTPException(status_code=400, detail="没有提供更新字段")

        cfg = state.model_config_mgr.update_config(agent_id, updates)
        return ApiResponse(
            message=f"Agent {agent_id} 配置已更新",
            data=cfg.to_dict()
        )

    @app.get("/api/model-providers")
    async def get_model_providers():
        """获取所有可用的模型提供商和模型列表"""
        from core.llm_adapter import get_available_providers
        return ApiResponse(data=get_available_providers())

    # ── 流水线执行状态 API ───────────────────────────
    @app.get("/api/pipelines/{pipeline_id}/status")
    async def get_pipeline_exec_status(pipeline_id: str):
        """获取流水线执行状态（含异步执行进度）"""
        state = get_state()

        # 先查异步执行状态
        running_info = state._running_pipelines.get(pipeline_id)

        # 查流水线实例
        pipeline = state.orchestrator.pipeline_engine.get_pipeline(pipeline_id)
        if not pipeline:
            raise HTTPException(status_code=404, detail="流水线不存在")

        result = pipeline.to_dict()
        if running_info:
            result["execution_status"] = running_info
        return ApiResponse(data=result)

    # ── 健康检查 ─────────────────────────────────────
    @app.get("/api/health")
    async def health_check():
        """健康检查"""
        return ApiResponse(data={
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "version": "1.0.0"
        })

    # ── SPA 入口 & 路由回退（必须放在所有 API 路由之后）──
    @app.get("/favicon.svg")
    async def favicon():
        """Favicon"""
        fav = frontend_dist / "favicon.svg"
        if fav.exists():
            return FileResponse(str(fav), media_type="image/svg+xml")
        raise HTTPException(status_code=404)

    @app.get("/{full_path:path}", response_class=HTMLResponse)
    async def spa_fallback(full_path: str):
        """SPA 路由回退 — 所有非 /api 请求返回 index.html"""
        if full_path.startswith("api"):
            raise HTTPException(status_code=404, detail="API endpoint not found")

        index_file = frontend_dist / "index.html"
        if index_file.exists():
            return index_file.read_text(encoding='utf-8')

        legacy_index = legacy_static / "index.html"
        if legacy_index.exists():
            return legacy_index.read_text(encoding='utf-8')

        return HTMLResponse(
            "<h1>OpenClaw Dashboard</h1>"
            "<p>前端未构建。请先运行 <code>cd frontend && npm install && npm run build</code></p>"
        )

    return app


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 直接运行
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

if __name__ == "__main__":
    import uvicorn
    app = create_app(str(PROJECT_ROOT))
    uvicorn.run(app, host="0.0.0.0", port=8080, reload=False)
